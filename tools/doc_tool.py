import PyPDF2
from PIL import Image
import docx
import pandas as pd
from pdf2image import convert_from_path
import re
import io

try:
    import pypdfium2 as pdfium
except Exception:
    pdfium = None


def read_pdf(file_path: str) -> str:
    text = ""

    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() or ""

    return text.strip()


def pdf_contains_images(file_path: str) -> bool:
    """
    Detect whether a PDF contains embedded image objects.
    """
    try:
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                resources = page.get("/Resources")
                if not resources:
                    continue

                xobject = resources.get("/XObject") if hasattr(resources, "get") else None
                if not xobject:
                    continue

                xobject = xobject.get_object()
                for obj in xobject.values():
                    resolved = obj.get_object()
                    if resolved.get("/Subtype") == "/Image":
                        return True
    except Exception:
        return False

    return False


def detect_language(text: str) -> str:
    if not text or not text.strip():
        return "unknown"

    try:
        from langdetect import detect

        return detect(text)
    except Exception:
        normalized = text.lower()
        words = re.findall(r"\w+", normalized)
        if not words:
            return "unknown"

        english_stopwords = {
            "the",
            "and",
            "is",
            "in",
            "to",
            "of",
            "a",
            "that",
            "it",
            "for",
            "on",
            "with",
            "as",
            "this",
            "be",
            "are",
            "or",
            "was",
            "by",
            "an",
        }
        stopword_count = sum(1 for w in words if w in english_stopwords)
        if stopword_count / len(words) >= 0.15:
            return "english"
        if re.search(r"[а-яА-Я]", text):
            return "russian"
        if re.search(r"[\u4e00-\u9fff]", text):
            return "chinese"
        if re.search(r"[áéíóúñçüãõ]", normalized):
            return "spanish"
        if re.search(r"[àâæçéèêëîïôœùûüÿ]", normalized):
            return "french"
        return "unknown"


def read_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])


def extract_pdf_pages_as_images(file_path: str, max_pages: int | None = None) -> list[bytes]:
    """
    Convert PDF pages to PNG bytes for vision analysis.

    Strategy:
    - Try high DPI first (300), then lower fallback DPIs for problematic files.
    - Convert page-by-page to reduce memory pressure on large PDFs.
    - Try with and without pdftocairo backend for compatibility.
    """
    images = []

    try:
        with open(file_path, "rb") as f:
            page_count = len(PyPDF2.PdfReader(f).pages)
    except Exception:
        page_count = 0

    if page_count == 0:
        return images

    last_page = min(page_count, max_pages) if max_pages is not None else page_count

    dpis = [300, 220, 150]

    # Strategy 1 (preferred): pypdfium2 (Poppler-free)
    if pdfium is not None:
        for dpi in dpis:
            try:
                doc = pdfium.PdfDocument(file_path)
                images = []
                for page_idx in range(last_page):
                    page = doc[page_idx]
                    bitmap = page.render(scale=dpi / 72)
                    pil_image = bitmap.to_pil()
                    buffer = io.BytesIO()
                    pil_image.save(buffer, format="PNG")
                    images.append(buffer.getvalue())

                if images:
                    return images
            except Exception:
                continue

    # Strategy 2: pdf2image + poppler (if available)
    backends = [True, False]  # use_pdftocairo

    for dpi in dpis:
        images = []
        converted_count = 0

        for page_num in range(1, last_page + 1):
            page_done = False

            for use_cairo in backends:
                try:
                    page_images = convert_from_path(
                        file_path,
                        dpi=dpi,
                        first_page=page_num,
                        last_page=page_num,
                        fmt="png",
                        use_pdftocairo=use_cairo,
                    )

                    if not page_images:
                        continue

                    buffer = io.BytesIO()
                    page_images[0].save(buffer, format="PNG")
                    images.append(buffer.getvalue())
                    converted_count += 1
                    page_done = True
                    break
                except Exception:
                    continue

            if not page_done:
                continue

        if converted_count > 0 and images:
            return images

    return []


def read_spreadsheet(file_path: str) -> dict:
    try:
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)

        columns = list(df.columns.astype(str))
        row_count = len(df)
        preview = df.head(5).to_csv(index=False)

        return {
            "status": "success",
            "type": "spreadsheet",
            "content": preview,
            "metadata": {
                "rows": row_count,
                "columns": len(columns),
                "column_names": columns[:10]
            }
        }
    except Exception as e:
        return {
            "status": "error",
            "message": str(e)
        }


def process_image(file_path: str) -> dict:
    try:
        # Validate that the image is readable, but do not run OCR.
        with Image.open(file_path) as img:
            img.verify()

        return {
            "status": "success",
            "type": "image",
            "content": "Image loaded successfully. Vision model can analyze visual content and text directly from the image."
        }

    except Exception as e:
        return {
            "status": "error",
            "message": str(e)
        }


def process_document(file_path: str) -> dict:
    try:
        lower_path = file_path.lower()

        if lower_path.endswith(".pdf"):
            image_bytes_list = extract_pdf_pages_as_images(file_path)
            if image_bytes_list:
                return {
                    "status": "success",
                    "type": "pdf_image",
                    "content": "PDF pages were converted to high-DPI images (300 DPI) for direct Gemma Vision analysis.",
                    "language": "unknown",
                    "image_bytes_list": image_bytes_list
                }

            return {
                "status": "error",
                "message": (
                    "Unable to convert PDF pages to images for vision analysis. "
                    "No page image could be rendered. Please check poppler/pdf rendering support."
                )
            }

        elif lower_path.endswith(".docx"):
            text = read_docx(file_path)
            return {
                "status": "success",
                "type": "docx",
                "content": text,
                "language": detect_language(text)
            }

        elif lower_path.endswith((".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tif", ".tiff")):
            return process_image(file_path)
        elif lower_path.endswith(('.csv', '.xlsx', '.xls')):
            return read_spreadsheet(file_path)
        else:
            return {
                "status": "error",
                "message": "Unsupported file type"
            }

    except Exception as e:
        return {
            "status": "error",
            "message": str(e)
        }